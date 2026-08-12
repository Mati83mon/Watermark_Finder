"""Plain-text extraction from uploaded documents.

Extraction runs here rather than in the Worker because parsing PDF and DOCX at
the edge is neither cheap nor safe within the free CPU budget. The Worker
streams the R2 object to this service and receives text back.

PDF and DOCX support depend on ``pypdf`` and ``python-docx``. Both are declared
in ``requirements.txt``; if either is missing the extractor reports the format
as unsupported instead of failing the whole request.
"""

from __future__ import annotations

import io
import json
import logging
import re
from dataclasses import dataclass
from html.parser import HTMLParser

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".tsv", ".log", ".rst"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx", ".html", ".htm", ".json"}


class ExtractionError(ValueError):
    """Raised when a document cannot be turned into text."""


@dataclass
class ExtractionResult:
    text: str
    format: str
    pages: int | None = None
    truncated: bool = False
    notes: list[str] = None  # type: ignore[assignment]

    def as_dict(self) -> dict[str, object]:
        return {
            "text": self.text,
            "format": self.format,
            "pages": self.pages,
            "truncated": self.truncated,
            "notes": self.notes or [],
            "chars": len(self.text),
        }


class _TextHTMLParser(HTMLParser):
    _SKIP = {"script", "style", "head", "noscript"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1
        elif tag in {"p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skip_depth > 0:
            self._skip_depth -= 1
        elif tag in {"p", "div", "li", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self.parts.append(data)

    def text(self) -> str:
        joined = "".join(self.parts)
        return re.sub(r"\n{3,}", "\n\n", joined).strip()


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1250", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extension(filename: str) -> str:
    lowered = filename.lower().strip()
    dot = lowered.rfind(".")
    return lowered[dot:] if dot != -1 else ""


def _extract_pdf(data: bytes) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on install
        raise ExtractionError("PDF support requires the 'pypdf' package") from exc

    reader = PdfReader(io.BytesIO(data))
    notes: list[str] = []
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError("The PDF is encrypted and cannot be read") from exc
        notes.append("The PDF was encrypted with an empty password.")

    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:  # pragma: no cover - malformed page
            notes.append("At least one page could not be parsed and was skipped.")
    text = "\n\n".join(chunk.strip() for chunk in chunks if chunk.strip())
    if not text:
        notes.append(
            "No text layer was found. The file is probably a scan; OCR is out of scope here."
        )
    return ExtractionResult(text=text, format="pdf", pages=len(reader.pages), notes=notes)


def _extract_docx(data: bytes) -> ExtractionResult:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - depends on install
        raise ExtractionError("DOCX support requires the 'python-docx' package") from exc

    document = docx.Document(io.BytesIO(data))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(block for block in blocks if block.strip())
    return ExtractionResult(text=text, format="docx", notes=[])


def _extract_json(data: bytes) -> ExtractionResult:
    payload = json.loads(_decode(data))
    collected: list[str] = []

    def walk(node) -> None:
        if isinstance(node, str):
            collected.append(node)
        elif isinstance(node, dict):
            for value in node.values():
                walk(value)
        elif isinstance(node, list):
            for value in node:
                walk(value)

    walk(payload)
    return ExtractionResult(
        text="\n".join(collected),
        format="json",
        notes=["Every string value in the document was concatenated."],
    )


def extract(data: bytes, filename: str, max_chars: int = 200_000) -> ExtractionResult:
    """Extract plain text from ``data``, dispatching on the file extension."""
    if not data:
        raise ExtractionError("The uploaded file is empty")

    extension = _extension(filename)
    if extension == ".pdf" or data[:5] == b"%PDF-":
        result = _extract_pdf(data)
    elif extension == ".docx" or (data[:2] == b"PK" and extension in {"", ".zip"}):
        result = _extract_docx(data)
    elif extension in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(_decode(data))
        result = ExtractionResult(text=parser.text(), format="html", notes=[])
    elif extension == ".json":
        result = _extract_json(data)
    elif extension in TEXT_EXTENSIONS or extension == "":
        result = ExtractionResult(text=_decode(data), format="text", notes=[])
    else:
        raise ExtractionError(
            f"Unsupported file type '{extension}'. Supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    if len(result.text) > max_chars:
        result.text = result.text[:max_chars]
        result.truncated = True
        result.notes = (result.notes or []) + [
            f"Document truncated to the first {max_chars} characters."
        ]
    if not result.text.strip():
        raise ExtractionError("No readable text could be extracted from the file")
    return result


__all__ = [
    "ExtractionError",
    "ExtractionResult",
    "SUPPORTED_EXTENSIONS",
    "extract",
]
